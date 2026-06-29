import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

def generate_manuals():
    title = "Manual de Conexión de Binance a NeuroTrade"
    
    sections = [
        ("Introducción", "Este manual te guiará paso a paso para conectar tu cuenta de Binance al sistema automatizado de NeuroTrade, tanto en modo de prueba (Demo/Testnet) como en modo de operaciones con dinero real."),
        
        ("1. Crear tu cuenta y llaves en Binance",
         "1. Inicia sesión en tu cuenta de Binance desde un navegador web.\n"
         "2. Ve a tu perfil y selecciona 'Gestión de API' (API Management).\n"
         "3. Haz clic en 'Crear API' (Generada por el sistema).\n"
         "4. Ponle un nombre (ej. 'NeuroTrade Bot').\n"
         "5. Completa las verificaciones de seguridad.\n"
         "6. IMPORTANTE: En los permisos de la API, asegúrate de marcar 'Habilitar Spot & Margin Trading'. NO marques retiros.\n"
         "7. Copia tu 'API Key' y tu 'Secret Key'. (Atención: La Secret Key solo se muestra una vez, guárdala en un lugar seguro)."),
         
        ("2. Configurar Binance en NeuroTrade",
         "1. Abre la interfaz de NeuroTrade en tu computadora.\n"
         "2. En el menú lateral izquierdo, haz clic en 'Configuración' (o el ícono del engranaje).\n"
         "3. Baja hasta la sección de credenciales de broker.\n"
         "4. Selecciona 'Binance' en el menú desplegable de 'Broker'.\n"
         "5. En el campo de 'Correo', pega tu API Key de Binance.\n"
         "6. En el campo de 'Contraseña', pega tu Secret Key de Binance."),
         
        ("3. ¿Demo o Real?",
         "NeuroTrade tiene un interruptor para elegir el tipo de cuenta:\n\n"
         "- CUENTA REAL: Al seleccionar 'Real', NeuroTrade utilizará tus llaves de API normales para operar en el mercado real de Binance usando tu dinero de la billetera Spot. Asegúrate de tener saldo en USDT para operar.\n\n"
         "- CUENTA DEMO (Cuidado): En Binance, a diferencia de otros brokers, el entorno 'Demo' se conecta a una red de desarrolladores llamada 'Binance Testnet'. Tus llaves normales de Binance NO funcionan en la Testnet. Si conectas tus llaves normales en modo 'Demo', el bot leerá saldo de $0.00 y no abrirá operaciones, además de que los gráficos se mostrarán como una línea plana. Si quieres usar la cuenta Demo en Binance, deberás generar llaves especiales en testnet.binance.vision.\n\n"
         "Recomendación para pruebas: Si quieres probar el bot sin dinero real usando el mercado vivo, te recomendamos usar IQ Option en modo Demo."),
         
        ("4. Ajustar el Riesgo (¡Muy Importante!)",
         "Antes de encender el motor en cuenta Real, ve a la pestaña 'Configuración' de NeuroTrade y baja a la sección de Riesgo. El 'Monto de Operación' determina de cuánto será cada compra que haga el bot. Te recomendamos bajar este monto a $5 o $10 dólares para tus primeras pruebas en Real, para que compruebes la efectividad de la Inteligencia Artificial de forma segura."),
         
        ("5. ¡Encender el Motor!",
         "1. Guarda los cambios haciendo clic en 'Guardar' o 'Conectar'.\n"
         "2. Ve a la pantalla principal ('Monitor V7').\n"
         "3. Haz clic en el botón 'ENCENDER MOTOR'.\n"
         "4. Verás los eventos del sistema empezar a analizar el mercado. Cuando el saldo aparezca correctamente y las velas del gráfico comiencen a moverse, ¡estarás listo para generar ganancias!")
    ]

    # Generate DOCX
    docx_path = r"C:\Users\DOCENTE\Desktop\Proyectos anti\NeuroTrade\Manual_Binance_NeuroTrade.docx"
    doc = Document()
    doc.add_heading(title, 0)
    
    for heading, text in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(text)
        
    doc.save(docx_path)

    # Generate PDF
    pdf_path = r"C:\Users\DOCENTE\Desktop\Proyectos anti\NeuroTrade\Manual_Binance_NeuroTrade.pdf"
    pdf = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    title_style = styles['Heading1']
    title_style.alignment = 1 # Center
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))
    
    body_style = styles['Normal']
    body_style.spaceAfter = 12
    
    heading_style = styles['Heading2']
    heading_style.spaceBefore = 12
    heading_style.spaceAfter = 6
    
    for heading, text in sections:
        story.append(Paragraph(heading, heading_style))
        # Replace newlines with <br/> for ReportLab
        formatted_text = text.replace('\n', '<br/>')
        story.append(Paragraph(formatted_text, body_style))
        
    pdf.build(story)
    
    print(f"Archivos creados exitosamente:\n1. {docx_path}\n2. {pdf_path}")

if __name__ == '__main__':
    generate_manuals()
